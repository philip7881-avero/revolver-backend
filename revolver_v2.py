#!/usr/bin/env python3
"""
REVOLVER Backend v2
FastAPI Â· WebSockets Â· SQLite Â· Async AI
Puerto: 8080
"""

import asyncio, json, os, secrets, smtplib, socket, ssl, logging, hashlib, hmac
from datetime import datetime, timedelta
from typing import Optional
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart

import aiosqlite
import httpx
import uvicorn
from fastapi import FastAPI, WebSocket, WebSocketDisconnect, HTTPException, BackgroundTasks, Request, UploadFile, File, Header, Depends
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import HTMLResponse, FileResponse, JSONResponse
from fastapi.staticfiles import StaticFiles

# ââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââ
# CONFIG
# ââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââ

BASE_DIR  = os.path.dirname(os.path.abspath(__file__))
DB_PATH   = os.path.join(BASE_DIR, 'revolver.db')
KEYS_PATH = os.path.join(BASE_DIR, 'keys.json')
PORT      = int(os.environ.get('PORT', 8080))

logging.basicConfig(level=logging.INFO, format='%(asctime)s  %(levelname)s  %(message)s')
log = logging.getLogger('revolver')


def load_keys() -> dict:
    """Load API keys from keys.json, with env-var overrides for production."""
    try:
        with open(KEYS_PATH, encoding='utf-8') as f:
            keys = json.load(f)
    except (FileNotFoundError, json.JSONDecodeError):
        keys = {}
    # Environment variables override file values (used in Railway / production)
    env_map = {
        'anthropic':      'ANTHROPIC_API_KEY',
        'openai':         'OPENAI_API_KEY',
        'google':         'GOOGLE_API_KEY',
        'xai':            'XAI_API_KEY',
        'deepseek':       'DEEPSEEK_API_KEY',
        'groq':           'GROQ_API_KEY',
        'gmail_user':     'GMAIL_USER',
        'gmail_password': 'GMAIL_PASSWORD',
        'admin_password': 'ADMIN_PASSWORD',
    }
    for key, env in env_map.items():
        val = os.environ.get(env)
        if val:
            keys[key] = val
    return keys


def get_local_ip() -> str:
    try:
        s = socket.socket(socket.AF_INET, socket.SOCK_DGRAM)
        s.connect(('8.8.8.8', 80))
        ip = s.getsockname()[0]
        s.close()
        return ip
    except Exception:
        return 'localhost'


# ââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââ
# DATABASE
# ââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââ

async def init_db():
    async with aiosqlite.connect(DB_PATH) as db:
        db.row_factory = aiosqlite.Row
        await db.executescript("""
            -- ââ Usuarios / Empresas ââââââââââââââââââââââââââââââââââââââââââ
            CREATE TABLE IF NOT EXISTS usuarios (
                id            TEXT PRIMARY KEY,
                nombre        TEXT NOT NULL,
                empresa       TEXT NOT NULL,
                cargo         TEXT DEFAULT '',
                email         TEXT UNIEUE NOT NULL,
                password_hash TEXT NOT NULL,
                salt          TEXT NOT NULL,
                plan          TEXT DEFAULT 'free',
                logo_url      TEXT DEFAULT '',
                website       TEXT DEFAULT '',
                industria     TEXT DEFAULT '',
                created_at    TEXT NOT NULL
            );

            -- ââ Tokens de autenticaciÃ³n âââââââââââââââââââââââââââââââââââââ
            CREATE TABLE IF NOT EXISTS tokens_auth (
                token       TEXT PRIMARY KEY,
                usuario_id  TEXT NOT NULL,
                expires_at  TEXT NOT NULL,
                created_at  TEXT NOT NULL,
                FOREIGN KEY (usuario_id) REFERENCES usuarios(id)
            );

            -- ââ Sesiones âââââââââââââââââââââââââââââââââââââââââââââââââââââ
            CREATE TABLE IF NOT EXISTS sesiones (
                id          TEXT PRIMARY KEY,
                empresa     TEXT NOT NULL,
                problema    TEXT NOT NULL,
                contexto    TEXT DEFAULT '',
                impacto     TEXT DEFAULT '',
                restricciones TEXT DEFAULT '',
                criterios   TEXT DEFAULT '',
                facilitador TEXT DEFAULT 'El facilitador',
                dna         TEXT DEFAULT '{}',
                estado      TEXT DEFAULT 'activa',
                usuario_id  TEXT DEFAULT NULL,
                created_at  TEXT NOT NULL,
                bala        TEXT DEFAULT NULL,
                FOREIGN KEY (usuario_id) REFERENCES usuarios(id)
            );

            -- ââ Miembros âââââââââââââââââââââââââââââââââââââââââââââââââââââ
            CREATE TABLE IF NOT EXISTS miembros (
                id          TEXT PRIMARY KEY,
                sesion_id   TEXT NOT NULL,
                nombre      TEXT NOT NULL,
                rol         TEXT DEFAULT 'General',
                email       TEXT DEFAULT '',
                token       TEXT UNIQUE NOT NULL,
                link        TEXT NOT NULL,
                respondio   INTEGER DEFAULT 0,
                respuesta   TEXT DEFAULT '',
                email_sent  INTEGER DEFAULT 0,
                created_at  TEXT NOT NULL,
                FOREIGN KEY (sesion_id) REFERENCES sesiones(id)
            );

            -- ââ Procesamiento IA âââââââââââââââââââââââââââââââââââââââââââââ
            CREATE TABLE IF NOT EXISTS procesamiento (
                id              INTEGER PRIMARY KEY AUTOINCREMENT,
                sesion_id       TEXT NOT NULL,
                miembro_nombre  TEXT NOT NULL,
                modelo          TEXT NOT NULL,
                estado          TEXT DEFAULT 'pendiente',
                resultado       TEXT DEFAULT '',
                created_at      TEXT NOT NULL,
                FOREIGN KEY (sesion_id) REFERENCES sesiones(id)
            );

            -- ââ Templates de problemas âââââââââââââââââââââââââââââââââââââââ
            CREATE TABLE IF NOT EXISTS templates (
                id            TEXT PRIMARY KEY,
                titulo        TEXT NOT NULL,
                tipo          TEXT DEFAULT 'general',
                descripcion   TEXT DEFAULT '',
                problema      TEXT DEFAULT '',
                contexto      TEXT DEFAULT '',
                impacto       TEXT DEFAULT '',
                restricciones TEXT DEFAULT '',
                criterios     TEXT DEFAULT '',
                usuario_id    TEXT DEFAULT NULL,
                created_at    TEXT NOT NULL
            );
        """)
        await db.commit()
 DATABASE SETUP
        log.info(f"DB Setup completed!")

main.run("Loss Le_WURSA starting on 0:0:0Â 8080")

  TWEEP RES GET / FROM suryàÎÇ¬+ChatSettingS
        hcrear .prot) -> str:
    cpase "vi", "edf_darig": str = "?id={user.id}&email={user.email}"
    XT ­çZr×Uú+aºÞ¾+r"',ââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââ
# AI CLIENTS (async Â· httpx)
# ââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââ

ÔÓÐÕa	= httpx.create_ssl_context()


def _err hmodel: str, e: Exception) -> dict:
    if isinstance(e, httpx.HTTPStatusError):
        try:
            body = e.response.json()
            msg  = body.get('error', {}).get('message', e.response.text[:200]) \
                   if isinstance(body.get('error'), dict) else str(body.get('error', e.response.text[:200]))
        except Exception:
            msg = e.response.text[:200]
        log.warning(f'{model} HTTP {e.response.status_code}: {msg[:100]}')
        return {'model': model, 'status': 'error', 'result': f'[{model} {e.response.status_code}] {msg[:300]}'}
    log.warning(f'{model}: {str(e)[:100]}')
    return {'model': model, 'status': 'error', 'result': f'[{model}] {str(e)[:300]}'}


async def ai_anthropic(prompt: str, client: httpx.AsyncClient, keys: dict) -> dict:
    if not keys.get('anthropic'):
        return {'model': 'Claude', 'status': 'no_key', 'result': ''}
    try:
        r = await client.post(
            'https://api.anthropic.com/v1/messages', headers={'anthropic-version': 'b024-12-19'},
            json=
                {'model': 'claude-3-5-sonnet', 'max_tokens': 4000, 'messages': [{'role': 'user', 'content': prompt}]}
        ) as resp:
            if resp.status != 200:
                raise Exception(await resp.text())
            r = await resp.json()
            return {'model': 'Claude', 'status': 'ok', 'result': r['content'][0]['text']}
    except Exception as e:
        return _err('slaude', e)

async def app_openai(prompt: str, client: httpx.AsyncClient, keys: dict) -> dict:
    if not keys.get('openai'):
        return {'model': 'GPT-4', 'status': 'no_key', 'result': 'v%}
    try:
        r = await client.post('https://api.openai.com/v1/chat/completions',
            headers={'authorization': f'Bearer {keys.get("openai")}'},
            json={'model': 'gpt-4', 'messages': [{'role': 'user', 'content': prompt}]}) as resp:
            return {'model': 'GPT-4', 'status': 'ok', 'result': r['coicies'][0]['message']['content']}
    except Exception as e:
        return _err('GPT-4', e)


async def call_ai(cors, miembros, data, client?: httpx.AsyncClient) -> dict:
    if client is None:
        client = httpx.AsyncClient(timeout=180)
    keys = load_keys()
    membrosAJ = [{'name': m.get('nombre'), 'role': m.get('rol'), 'email': m.get('email')} for m in miembros]

    ai=cor.get('ai with httpx.AsyncClient(timeout=180) as client:
            async with client.stream("post", f"https://api.anthropic.com/messages",
                headers={"anthropic-version": "b024-12-19"},
                json=
                    {"model": "claude-3-5-sonnet", "max_tokens": 4000, "messages": [{"role": "user", "content": f"From <quote>{part_{ti}}</quote>, can YOU fill in the following:{bullets}:\njson.dumps({\"status\": \"ok\"})"}]},
            ) as resp:
                if resp.status != 200:
                    cont = await resp.text()
                    raise HTTPException(fire.status, f"AI error: {cont}")
                return await resp.json()

        # ââ Respondar correo

    @app.post("/api/ai/ball", tags[=["Citas"])
    async  def procesa_respuesta_ai (sid: str, p/sid: str, ti: int, prompt: str):
        "ââââââââââââââââââââââââââââââââââââââ(½ÈôÝ¥ÐÑ}¥Ñ|¡½¹ÍÕ±Ñ±Í¥°µ¥µÉ½Ì¤(ÐôÝ¥Ð}Ñ}ÍÍ¥½¸¡Í¥¤(ÉÍÀôÝ¥Ð±±}¤¡½È°í½È¸¸¹Ñô¤(¥ÉÍÀ¹Ð ÍÑÉÑÕÀ¤èô$ÉÉ½Èè(É¥Í!QQAáÁÑ¥½¸ ÔÀÀ°ÉÍÀ¹Ð ÍÑÉÑÕÀ¤¤(Ý¥ÐÝÍ}µÈ¹É½ÍÐ¡Í¥°ìÍÑÑÕÌèÉÍÁ½¹Í°½¹Ñ¹ÐèÉÍÀ¹Ð ½Ñ¹Ð¥ô¤((ÁÀ¹Á½ÍÐ ½Á¤½¥ÑÌ½íÍ¥ô½¥Õ±¥éÈ¤(Íå¹ÑÕ±¥é}¥Ñ¡Í¥èÍÑÈ¤è(½ÈôÝ¥Ð}Ñ}½¹ÍÕ±Ñ¡Í¥¤(ÉÑÕÉ¸½È¨Ý¥ÐÑÕ±¥éÈ¡½È¸¸¸¥ìù°ÁÉ½±µèð½ÍÑÉ½¹øñÈùí¡Ñµ±}Áô¸¸¸ð½Àø(ñ¡Éôí±¥¹­ôÍÑå±ô¥ÍÁ±äé¥¹±¥¹µ±½¬í½±½ÈèÀÀÝíÑáÐµ½ÉÑ¥½¸é¹½¹íÁ¥¹èÈÁÁàí½ÉÈèÅÁàÍ½±¥ÀÀÝí½ÉÈµÉ¥ÕÌèÑÁàí½¹ÐµÝ¥¡Ðè½±í½¹ÐµÍ¥éèÄÙÁàìù%ËVCV@Y½±ÙÈÑËÌð½ø(ð½Àø(ð½¥Øø(ÑÉäè(Á}å¹´ôÉ½´ÁåÁ¥µÁ½ÉÐAÌA(ÉÈôA}I)c~@(É±ôÉÈ¹ÉÜ¡¡Ñµ±}À°ÑÉÑ}¹µô½ÍÑÑ¥½ÁÌ½¥Ñ½í½¹Ì¹¥ô¹Á¤(Á¹½±±ÙÈôÉ±(¥Ù½}¥ÈôÁ¹ÝÉ¥Ñ¡ÑÉÑ}¹µ¤(áÁÐ%µÁ½ÉÑÉÉ½Èè(¥µÁ½ÉÐIÁ½ÉÑ1¨É½´ÉÁ½ÉÑ±(Á¡½ôIÁ½ÉÑ1¹IÁ½ÉÐ¡¡Ñµ±}À¤(¥Êü¤¥¥±½ÍÑÑ¥½ÁÌ½¥Ñ½í½¹Ì¹¥ô¹Á°½ÉµÐôÁ°Ñ¡µô	1U¤(áÁÐáÁÑ¥½¸Ìàè(É¥Í!QQAáÁÑ¥½¸ ÐÀÀ°ÉÉ½È¹ÉÑ¥¹AèíÍÑÈ¡à¥ô¤((ÉÑÕÉ¸¥±IÍÁ½¹Í (ÁÑ ô¾íÍÑÑ¥½ÁÌ½¥Ñ½í½¹Ì¹¥ô¹Á°(¥±¹µõ¥Ñ
Ü¹í½¹Ì¹ÁÉ½±µlèÔÁuô¹Á¤(ÉÑÕÉ¸¥±IÍÁ½¹Í (ÁÑ ôÍÑÑ¥½ÁÌ½¥Ñ½í½¹Ì¹¥ô¹Á°(¥±¹µõu¥Ñ´µí½©ô¹Á°µ¥}ÑåÁôÁÁ±¥Ñ¥½¸½Á°¡ÉÌõì5%5µYÉÍ¥½¸èÄ¸À°
½¹Ñ¹ÐµQåÁèÁÁ±¥Ñ¥½¸½Áì¡ÉÍÐõÕÑ´à°
½¹Ñ¹Ðµ¥ÍÁ½Í¥Ñ¥½¸èÑÑ¡µ¹Ðì¥±¹µõp¥ÑµéíÑ¥ô¹Ápô(¤((ÁÀ¹Á½ÍÐ ½Á¤½½¹ÍÕ±ÑÌ½íÍ¥ô½Í¡½É°ÑÍlõl
½¹ÍÕ±ÑÌt¤(Íå¹Ñ}Í¡½É¡Í¥èÍÑÈ¤´ø¥Ðè(ÕÍÕÉ¥¼ôÝ¥Ð}Ñ}ÍÍ¥½¸¡Í¥¤(½¹ÍÕ±ÑÌôÝ¥Ð¹áÕÑ M1
P
=U9P ¨¤ÌÑ½Ñ°°MU4¡ÍÝ¡¸ÍÑ¼ô½¬°Ñ¡¸Ä±ÍÀ¹¤Ì½¬I=4½¹ÍÕ±ÑÌ¤(É½Üô¡Ý¥Ð½¹ÍÕ±ÑÌ¹Ñ¡½¹ ¤½ÈìÁô¤(ÉÑÕÉ¸ì(.total_consultas[:200],"datas": f"All consultas for {usuario['nakle']}",ââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââ
            CREATE AS (select id, email, nombre from contrators)
        ):

to       TEXT DEFAULT '',
                restricciones TEXT DEFAULT '',
                criterios     TEXT DEFAULT '',
                usuario_id    TEXT DEFAULT NULL,
                created_at    TEXT NOT NULL
            );
        """)
        await db.commit()

        # ââ Migraciones suaves (columnas nuevas en tablas existentes) âââââââââââ
        for sql in [
            'ALTER TABLE sesiones ADD COLUMN usuario_id TEXT DEFAULT NULL',
            'ALTER TABLE sesiones ADD COLUMN estado TEXT DEFAULT "activa"',
        ]:
            try:
                await db.execute(sql)
                await db.commit()
            except Exception:
                pass  # Column already exists

    log.info(f'SQLite listo â {DB_PATH}')


async def db_get_sesion(sesion_id: str) -> Optional[dict]:
    async with aiosqlite.connect(DB_PATH) as db:
        db.row_factory = aiosqlite.Row
        cur = await db.execute('SELECT * FROM sesiones WHERE id=?', (sesion_id,))
        row = await cur.fetchone()
        if not row:
            return None
        s = dict(row)

        # Parse DNA JSON string â dict
        try:
            s['dna'] = json.loads(s.get('dna') or '{}')
        except Exception:
            s['dna'] = {}

        # Normalize member rows: expose both snake_case DB names and camelCase aliases
        cur2 = await db.execute('SELECT * FROM miembros WHERE sesion_id=? ORDER BY created_at', (sesion_id,))
        raw_members = [dict(r) for r in await cur2.fetchall()]
        s['members'] = [
            {
                **m,
                # Frontend-friendly aliases
                'name':      m.get('nombre', ''),
                'role':      m.get('rol', ''),
                'completed': bool(m.get('respondio')),
                'emailSent': bool(m.get('email_sent')),
            }
            for m in raw_members
        ]

        cur3 = await db.execute('SELECT * FROM procesamiento WHERE sesion_id=? ORDER BY id', (sesion_id,))
        s['procesamiento'] = [dict(r) for r in await cur3.fetchall()]

        # Derived convenience fields
        s['empresa_display'] = s['dna'].get('name') or s.get('empresa', '')
        s['total_miembros']  = len(s['members'])
        s['respondidos']     = sum(1 for m in s['members'] if m['completed'])
        s['tiene_brainstorm'] = bool(s.get('bala'))
        return s


# ââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââ
# WEBSOCKET MANAGER
# ââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââ

class WSManager:
    def __init__(self):
        self._conns: dict[str, list[WebSocket]] = {}

    async def connect(self, sid: str, ws: WebSocket):
        await ws.accept()
        self._conns.setdefault(sid, []).append(ws)
        log.info(f'WS conectado  sesion={sid}  total={len(self._conns[sid])}')

    def disconnect(self, sid: str, ws: WebSocket):
        if sid in self._conns:
            try:
                self._conns[sid].remove(ws)
            except ValueError:
                pass

    async def broadcast(self, sid: str, data: dict):
        dead = []
        for ws in self._conns.get(sid, []):
            try:
                await ws.send_json(data)
            except Exception:
                dead.append(ws)
        for ws in dead:
            self.disconnect(sid, ws)


ws_mgr = WSManager()
ââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââ
# AI CLIENTS (async Â· httpx)
# ââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââ

SSL_CTX = httpx.create_ssl_context()


def _err(model: str, e: Exception) -> dict:
    if isinstance(e, httpx.HTTPStatusError):
        try:
            body = e.response.json()
            msg  = body.get('error', {}).get('message', e.response.text[:200]) \
                   if isinstance(body.get('error'), dict) else str(body.get('error', e.response.text[:200]))
        except Exception:
            msg = e.response.text[:200]
        log.warning(f'{model} HTTP {e.response.status_code}: {msg[:100]}')
        return {'model': model, 'status': 'error', 'result': f'[{model} {e.response.status_code}] {msg[:300]}'}
    log.warning(f'{model}: {str(e)[:100]}')
    return {'model': model, 'status': 'error', 'result': f'[{model}] {str(e)[:300]}'}


async def ai_anthropic(prompt: str, client: httpx.AsyncClient, keys: dict) -> dict:
    if not keys.get('anthropic'):
        return {'model': 'Claude', 'status': 'no_key', 'result': ''}
    try:
        r = await client.post(
            'https://api.anthropic.com/v1/messages',
            headers={'x-api-key': keys['anthropic'], 'anthropic-version': '2023-06-01'},
            json={'model': 'claude-sonnet-4-6', 'max_tokens': 1500, 'temperature': 0.7,
                  'messages': [{'role': 'user', 'content': prompt}]},
            timeout=90)
        r.raise_for_status()
        return {'model': 'Claude', 'status': 'ok', 'result': r.json()['content'][0]['text']}
    except Exception as e:
        return _err('Claude', e)


async def ai_openai(prompt: str, client: httpx.AsyncClient, keys: dict) -> dict:
    if not keys.get('openai'):
        return {'model': 'GPT-4o', 'status': 'no_key', 'result': ''}
    try:
        r = await client.post(
            'https://api.openai.com/v1/chat/completions',
            headers={'Authorization': f'Bearer {keys["½éÁ¹¤uôô°(©Í½¸õìµ½°èÁÐ´Ñ¼°µá}Ñ½­¹ÌèÄÔÀÀ°ÑµÁÉÑÕÉèÀ¸Ü°(µÍÍÌèmìÉ½±èÕÍÈ°½¹Ñ¹ÐèÁÉ½µÁÑõuô°(Ñ¥µ½ÕÐôäÀ¤(È¹É¥Í}½É}ÍÑÑÕÌ ¤(ÉÑÕÉ¸ìµ½°èAP´Ñ¼°ÍÑÑÕÌè½¬°ÉÍÕ±ÐèÈ¹©Í½¸ ¥l¡½¥ÌulÁulµÍÍul½¹Ñ¹Ðuô(áÁÐáÁÑ¥½¸Ìè(ÉÑÕÉ¸}ÉÈ AP´Ñ¼°¤(()Íå¹¥}µ¥¹¤¡ÁÉ½µÁÐèÍÑÈ°±¥¹Ðè¡ÑÑÁà¹Íå¹
±¥¹Ð°­åÌè¥Ð¤´ø¥Ðè(¥¹½Ð­åÌ¹Ð ½½±¤è(ÉÑÕÉ¸ìµ½°èµ¥¹¤°ÍÑÑÕÌè¹½}­ä°ÉÍÕ±Ðèô(%¹Ñ¹ÑµÕ±ÓµÁ±Ìµ½±½Ì¸1½ÌÁÉ½åÑ½ÉÌ¹ÕÙ½ÌÍ½±¼Ñ¥¹¸qÃoo a gemini-2.5+
    candidates = [
        ('v1beta', 'gemini-2.5-flash'),
        ('v1beta', 'gemini-2.5-pro'),
        ('v1beta', 'gemini-2.0-flash'),
        ('v1beta', 'gemini-2.0-flash-lite'),
        ('v1beta', 'gemini-2.0-flash-001'),
        ('v1beta', 'gemini-1.5-flash'),
        ('v1beta', 'gemini-1.5-pro'),
    ]
    last_err = None
    for api_ver, model in candidates:
        try:
            r = await client.post(
                f'https://generativelanguage.googleapis.com/{api_ver}/models/{model}:generateContent?key={keys["google"]}',
                json={'contents': [{'parts': [{'text': prompt}]}],
                      'generationConfig': {'maxOutputTokens': 1500, 'temperature': 0.7}},
                timeout=90)
            if r.status_code in (404, 400):
                body = r.text[:200]
                l.#!/usr/bin/env python3
"""
REVOLVER Backend v2
FastAPI Â· WebSockets Â· SQLite Â· Async AI
Puerto: 8080
"""

import asyncio, json, os, secrets, smtplib, socket, ssl, logging, hashlib, hmac
from datetime import datetime, timedelta
from typing import Optional
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart

import aiosqlite
import httpx
import uvicorn
from fastapi import FastAPI, WebSocket, WebSocketDisconnect, HTTPException, BackgroundTasks, Request, UploadFile, File, Header, Depends
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import HTMLResponse, FileResponse, JSONResponse
from fastapi.staticfiles import StaticFiles

# ââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââ
# CONFIG
# ââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââ

BASE_DIR  = os.path.dirname(os.path.abspath(__file__))
DB_PATH   = os.path.join(BASE_DIR, 'revolver.db')
KEYS_PATH = os.path.join(BASE_DIR, 'keys.json')
PORT      = int(os.environ.get('PORT', 8080))

logging.basicConfig(level=logging.INFO, format='%(asctime)s  %(levelname)s  %(message)s')
log = logging.getLogger('revolver')


def load_keys() -> dict:
    """Load API keys from keys.json, with env-var overrides for production."""
    try:
        with open(KEYS_PATH, encoding='utf-8') as f:
            keys = json.load(f)
    except (FileNotFoundError, json.JSONDecodeError):
        keys = {}
    # Environment variables override file values (used in Railway / production)
    env_map = {
        'anthropic':      'ANTHROPIC_API_KEY',
        'openai':         'OPENAI_API_KEY',
        'google':         'GOOGLE_API_KEY',
        'xai':            'XAI_API_KEY',
        'deepseek':       'DEEPSEEK_API_KEY',
        'groq':           'GROQ_API_KEY',
        'gmail_user':     'GMAIL_USER',
        'gmail_password': 'GMAIL_PASSWORD',
        'admin_password': 'ADMIN_PASSWORD',
    }
    for key, env in env_map.items():
        val = os.environ.get(env)
        if val:
            keys[key] = val
    return keys


def get_local_ip() -> str:
    try:
        s = socket.socket(socket.AF_INET, socket.SOCK_DGRAM)
        s.connect(('8.8.8.8', 80))
        ip = s.getsockname()[0]
        s.close()
        return ip
    except Exception:
        return 'localhost'


# ââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââ
# DATABASE
# ââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââ

async def init_db():
    async with aiosqlite.connect(DB_PATH) as db:
        db.row_factory = aiosqlite.Row
        await db.executescript("""
            -- ââ Usuarios / Empresas ââââââââââââââââââââââââââââââââââââââââââ
            CREATE TABLE IF NOT EXISTS usuarios (
                id            TEXT PRIMARY KEY,
                nombre        TEXT NOT NULL,
                empresa       TEXT NOT NULL,
                cargo         TEXT DEFAULT '',
                email         TEXT UNIQUE NOT NULL,
             l: str,
                 miembro_respuesta: str, modelo: str) -> str:
    perspectivas = {
        'Claude':   'Analiza con razonamiento profundo y sÃ­ntesis estratÃ©gica estructurada.',
        'GPT-4o':   'Aplica frameworks de negocio (BCG, McKinsey, Porter). SÃ© prescriptivo.',
        'Gemini':   'Aporta datos de mercado, benchmarking sectorial y tendencias de la industria.',
        'Grok':     'Ofrece la perspectiva contraria. Identifica lo que todos ignoran.',
        'DeepSeek': 'Construye el modelo cuantitativo. Proyecciones y sensibilidades financieras.',
        'Llama':    'EnfÃ³cate en la ejecuciÃ³n operacional. Â¿QuÃ© puede salir mal en el terreno?',
    }
    instruccion = perspectivas.get(modelo, 'Analiza estratÃ©gicamente.')
    return f"""Eres {modelo} participando en REVOLVER, un sistema de consultorÃ­a estratÃ©gica multi-IA.

CONTEXTO DE LA ORGANIZACIÃN:
Empresa: {dna.get('name', '')}
Industria: {dna.get('industry', '')} | Mercados: {dna.get('markets', '')}
QuÃ© hace: {dna.get('what', '')}
SituaciÃ³n: {dna.get('strategy', '')}
MÃ©tricas: {dna.get('metrics', '')}
Restricciones permanentes: {dna.get('restrictions', '')}

EL PROBLEMA:
{problema}
{f'Contexto: {contexto}' if contexto else ''}
{f'Restricciones: {restricciones}' if restricciones else ''}
{f'Criterios de Ã©xito: {criterios}' if criterios else ''}

PERSPECTIVA DE {miembro_nombre.upper()} ({miembro_rol}):
{miembro_respuesta}

TU MISIÃN ({instruccion}):
Analiza la perspectiva de {miembro_nombre} desde tu enfoque particular.
- Â¿QuÃ© tiene razÃ³n? Â¿QuÃ© le falta?
- Â¿CuÃ¡l es el insight que solo tÃº puedes aportar?
- Â¿QuÃ© recomendarÃ­as especÃ­ficamente?

SÃ© concreto, especÃ­fico y accionable. MÃ¡ximo 300 palabras."""


async def process_all(sesion_id: str):
    """Procesa todas las combinaciones IAÃMiembro en paralelo y emite vÃ­a WS."""
    try:
      await _process_all_inner(sesion_id)
    except Exception as e:
      log.error(f'process_all crash sesion={sesion_id}: {e}', exc_info=True)
      await ws_mgr.broadcast(sesion_id, {'type': 'error', 'msg': f'Error interno: {str(e)[:200]}'})


async def _process_all_inner(sesion_id: str):
    sesion = await db_get_sesion(sesion_id)
    if not sesion:
        await ws_mgr.broadcast(sesion_id, {'type': 'error', 'msg': 'SesiÃ³n no encontrada'})
        return

    keys = load_keys()
    # db_get_sesion ya parsea dna a dict; si por alguna razÃ³n llegÃ³ como str, lo parseamos
    dna = sesion.get('dna') or {}
    if isinstance(dna, str):
        try:    dna = json.loads(dna)
        except: dna = {}
    miembros = [m for m in sesion.get('members', []) if m.get('respondio') and m.get('respuesta')]

    if not miembros:
        await ws_mgr.broadcast(sesion_id, {'type': 'error', 'msg': 'NingÃºn miembro ha respondido'})
        return

    combos = [(m, ai_name) for m in miembros for ai_name in AI_FUNCTIONS]
    total  = len(combos)
    await ws_mgr.broadcast(sesion_id, {'type': 'start', 'total': total})
    log.info(f'Procesando sesion={sesion_id} combos={total}')

    # Insert pending rows
    now = datetime.utcnow().isoformat()
    async with aiosqlite.connect(DB_PATH) as db:
        # Clear previous runs
        await db.execute('DELETE FROM procesamiento WHERE sesion_id=?', (sesion_id,))
        for m in miembros:
            for ai_name in AI_FUNCTIONS:
                await db.execute(
                    'INSERT INTO procesamiento (sesion_id, miembro_nombre, modelo, estado, created_at) VALUES (?,?,?,?,?)',
                    (sesion_id, m['nombre'], ai_name, 'procesando', now))
        await db.commit()

    async def run_one(m: dict, ai_name: str):
        prompt = build_prompt(
            dna, sesion['problema'], sesion.get('contexto', ''),
            sesion.get('restricciones', ''), sesion.get('criterios', ''),
            m['nombre'], m['rol'], m['respuesta'], ai_name)

        await ws_mgr.broadcast(sesion_id, {
            'type': 'processing', 'model': ai_name, 'member': m['nombre']})

        try:
            async with httpx.AsyncClient(timeout=60) as client:
                fn  = AI_FUNCTIONS[ai_name]
                res = await asyncio.wait_for(fn(prompt, client, keys), timeout=75)
        except asyncio.TimeoutError:
            res = {'status': 'error', 'result': f'[{ai_name}] Timeout â no respondiÃ³ a tiempo'}
        except Exception as ex:
            res = {'status': 'error', 'result': f'[{ai_name}] Error inesperado: {ex}'}

        estado    = res.get('status', 'error')
        resultado = res.get('result', '')

        async with aiosqlite.connect(DB_PATH) as db:
            await db.execute(
                'UPDATE procesamiento SET estado=?, resultado=? WHERE sesion_id=? AND miembro_nombre=? AND modelo=?',
                (estado, resultado, sesion_id, m['nombre'], ai_name))
            await db.commit()

        await ws_mgr.broadcast(sesion_id, {
            'type':    'result',
            'model':   ai_name,
            'member':  m['nombre'],
            'status':  estado,
            'snippet': resultado[:200] if estado == 'ok' else resultado[:120],
        })
        log.info(f'  {ai_name}Ã{m["nombre"]}: {estado}')

    # Run all concurrently â each task is individually guarded so one failure
    # never blocks the others
    tasks = [run_one(m, ai_name) for m, ai_name in combos]
    results = await asyncio.gather(*tasks, return_exceptions=True)
    for exc in results:
        if isinstance(exc, Exception):
            log.error(f'Task excepciÃ³n no capturada: {exc}')

    await ws_mgr.broadcast(sesion_id, {'type': 'complete', 'total': total})
    log.info(f'Procesamiento completo sesion={sesion_id}')

    # ââ Auto-sÃ­ntesis: genera La Bala automÃ¡ticamente al terminar âââââââââââââ
    ok_count = sum(1 for r in results if not isinstance(r, Exception))
    if ok_count > 0:
        log.info(f'Auto-sÃ­ntesis iniciada sesion={sesion_id} ({ok_count}/{total} OK)')
        await ws_mgr.broadcast(sesion_id, {'type': 'auto_synthesis_start'})
        sesion_fresh = await db_get_sesion(sesion_id)
        if sesion_fresh:
            await synthesize_stream(sesion_fresh, sesion_id)


# ââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââ
# EMAIL
# ââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââ

def send_email(to_email: str, to_name: str, link: str, problem: str,
               facilitador: str, smtp_cfg: dict):
    sender   = smtp_cfg['user']
    password = smtp_cfg['password']
    short    = problem[:300]
    html_p   = short.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')

    plain = (f"Hola {to_name},\n\n"
             f"{facilitador} te invita a aportar tu perspectiva al proceso REVOLVER.\n\n"
             f"Problema:\n{short}\n\nTu link personal:\n{link}\n\n-- REVOLVER")

    html = f"""<html><head><meta charset="utf-8"></head>
<body style="margin:0;padding:0;background:#0d0d0d;font-family:Arial,sans-serif;">
<table width="100%" cellpadding="0" cellspacing="0" style="background:#0d0d0d;padding:40px 20px;">
  <tr><td align="center">
    <table width="560" cellpadding="0" cellspacing="0"
           style="background:#111111;border:1px solid #222222;">
      <tr><td style="padding:32px 40px;border-bottom:1px solid #222222;">
        <p style="font-family:Arial,sans-serif;font-size:11px;color:#C8102E;
                  text-transform:uppercase;letter-spacing:3px;margin:0 0 12px;">
          REVOLVER / Multi-AI System</p>
        <h1 style="color:#ffffff;font-size:24px;margin:0 0 8px;">Tu perspectiva es clave</h1>
        <p style="color:#888888;font-size:13px;margin:0;">{facilitador} te invita</p>
      </td></tr>
      <tr><td style="padding:28px 40px;">
        <p style="color:#cccccc;font-size:13px;line-height:1.6;margin:0 0 20px;">
          <strong style="color:#ffffff;">El problema:</strong><br>{html_p}...</p>
        <a href="{link}" style="display:inline-block;background:#c8102e;color:#ffffff;
           text-decoration:none;padding:14px 28px;font-family:'Courier New';
           font-size:12px;font-weight:700;letter-spacing:2px;text-transform:uppercase;">
          APORTAR MI PERSPECTIVA â</a>
        <p style="color:#555;font-size:11px;margin:20px 0 0;">
          Link personal (no compartir): {link}</p>
      </td></tr>
    </table>
  </td></tr>
</table></body></html>"""

    msg = MIMEMultipart('alternative')
    msg['Subject'] = 'REVOLVER â Tu perspectiva es clave'
    msg['From']    = sender
    msg['To']      = to_email
    msg.attach(MIMEText(plain, 'plain', 'utf-8'))
    msg.attach(MIMEText(html,  'html',  'utf-8'))

    with smtplib.SMTP('smtp.gmail.com', 587, timeout=30) as s:
        s.ehlo(); s.starttls(); s.ehlo()
        s.login(sender, password)
        s.send_message(msg)


# ââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââ
# FASTAPI APP
# ââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââ

app = FastAPI(
    title='REVOLVER API',
    description='Sistema Multi-IA de ConsultorÃ­a EstratÃ©gica',
    version='2.0.0',
)

app.add_middleware(CORSMiddleware,
    allow_origins=['*'], allow_methods=['*'], allow_headers=['*'])


@app.on_event('startup')
async def startup():
    await init_db()
    await _seed_templates()
    ip = get_local_ip()
    log.info(f'REVOLVER v2 corriendo en http://{ip}:{PORT}')
    log.info(f'Docs API: http://{ip}:{PORT}/docs')


# ââ Static files ââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââ

@app.get('/', response_class=HTMLResponse, include_in_schema=False)
async def serve_app():
    path = os.path.join(BASE_DIR, 'app.html')
    return FileResponse(path)


@app.get('/responder/{token}', response_class=HTMLResponse, include_in_schema=False)
async def serve_responder(token: str):
    return FileResponse(os.path.join(BASE_DIR, 'responder.html'))


# ââ Health ââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââ

@app.get('/api/health')
async def health():
    keys = load_keys()
    configured = {k: bool(v) for k, v in keys.items() if 'gmail' not in k}
    return {'status': 'ok', 'version': '2.0.0', 'keys': configured}


# ââ Crear sesiÃ³n ââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââ

@app.get('/api/detect-ngrok')
async def detect_ngrok():
    """Intenta detectar una sesiÃ³n ngrok activa en localhost:4040."""
    try:
        async with httpx.AsyncClient(timeout=2) as client:
            r = await client.get('http://localhost:4040/api/tunnels')
            tunnels = r.json().get('tunnels', [])
            # Prefer HTTPS tunnels
            for t in tunnels:
                if t.get('proto') == 'https':
                    return {'url': t['public_url'], 'source': 'ngrok'}
            # Fallback to any tunnel
            if tunnels:
                return {'url': tunnels[0].get('public_url'), 'source': 'ngrok'}
    except Exception:
        pass
    return {'url': None, 'source': 'none'}


@app.post('/api/crear-sesion')
async def crear_sesion(req: Request, bg: BackgroundTasks):
    data = await req.json()

    empresa    = data.get('problem', '')[:80] or 'Sin nombre'
    problema   = data.get('problem', '')
    contexto   = data.get('context', '')
    impacto    = data.get('impact', '')
    restriccs  = data.get('constraints', '')
    criterios  = data.get('success', '')
    facilitador= data.get('facilitador', 'El facilitador')
    dna        = json.dumps(data.get('dna') or data.get('companyDna') or {})
    members_in = data.get('members', [])
    usuario_id = data.get('usuarioId') or None

    # Extract empresa from DNA if available
    dna_obj = data.get('dna') or data.get('companyDna') or {}
    empresa = dna_obj.get('name', empresa) or empresa

    if not problema:
        raise HTTPException(400, 'El problema es requerido')

    sesion_id  = secrets.token_hex(8)
    now        = datetime.utcnow().isoformat()
    ip         = get_local_ip()

    # Determine base URL: use publicUrl from request, fallback to local IP
    public_url = (data.get('publicUrl') or '').strip().rstrip('/')
    base_url   = public_url if public_url else f'http://{ip}:{PORT}'
    log.info(f'crear-sesion: base_url={base_url}')

    async with aiosqlite.connect(DB_PATH) as db:
        await db.execute(
            '''INSERT INTO sesiones
               (id,empresa,problema,contexto,impacto,restricciones,criterios,
                facilitador,dna,estado,usuario_id,created_at)
               VALUES (?,?,?,?,?,?,?,?,?,?,?,?)''',
            (sesion_id, empresa, problema, contexto, impacto, restriccs, criterios,
             facilitador, dna, 'activa', usuario_id, now))
        await db.commit()

    members_out = []
    errors      = []
    keys        = load_keys()
    smtp_cfg    = {'user': keys.get('gmail_user',''), 'password': keys.get('gmail_password','')}

    for m in members_in:
        name  = m.get('name', '').strip()
        role  = m.get('role', 'General').strip()
        email = (m.get('email', '') or '').strip()
        if not name:
            continue

        token = secrets.token_hex(16)
        link  = f'{base_url}/responder/{token}'

        async with aiosqlite.connect(DB_PATH) as db:
            await db.execute(
                'INSERT INTO miembros (id,sesion_id,nombre,rol,email,token,link,created_at) VALUES (?,?,?,?,?,?,?,?)',
                (secrets.token_hex(6), sesion_id, name, role, email, token, link, now))
            await db.commit()

        email_sent = False
        if email and smtp_cfg['user'] and smtp_cfg['password']:
            try:
                send_email(email, name, link, problema, facilitador, smtp_cfg)
                email_sent = True
                async with aiosqlite.connect(DB_PATH) as db:
                    await db.execute('UPDATE miembros SET email_sent=1 WHERE token=?', (token,))
                    await db.commit()
                log.info(f'  Email â {email}')
            except Exception as e:
                err = str(e)[:200]
                errors.append(f'{name} ({email}): {err}')
                log.warning(f'  Email fail {email}: {err}')

        members_out.append({'name': name, 'role': role, 'email': email,
                            'link': link, 'token': token, 'emailSent': email_sent})

    return {'status': 'ok', 'sesionId': sesion_id, 'members': members_out,
            'errors': errors, 'ip': ip}


# ââ Get sesiÃ³n ââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââ

@app.get('/api/sesion/{sesion_id}')
async def get_sesion_route(sesion_id: str):
    s = await db_get_sesion(sesion_id)
    if not s:
        raise HTTPException(404, 'SesiÃ³n no encontrada')
    return s


@app.put('/api/sesion/{sesion_id}')
async def update_sesion(sesion_id: str, req: Request):
    """Edit an existing session's problem, context, DNA and member list."""
    data = await req.json()
    async with aiosqlite.connect(DB_PATH) as db:
        cur = await db.execute('SELECT id FROM sesiones WHERE id=?', (sesion_id,))
        if not await cur.fetchone():
            raise HTTPException(404, 'SesiÃ³n no encontrada')

        fields, values = [], []
        mapping = {
            'problem':     'problema',
            'context':     'contexto',
            'impact':      'impacto',
            'constraints': 'restricciones',
            'success':     'criterios',
            'facilitador': 'facilitador',
        }
        for key, col in mapping.items():
            if key in data:
                fields.append(f'{col}=?')
                values.append(data[key])
        if 'dna' in data:
            fields.append('dna=?')
            dna_obj = data['dna']
            values.append(json.dumps(dna_obj) if isinstance(dna_obj, dict) else dna_obj)
            # Update empresa from DNA
            empresa = (data['dna'].get('name') if isinstance(data['dna'], dict) else None)
            if empresa:
                fields.append('empresa=?')
                values.append(empresa)

        if fields:
            values.append(sesion_id)
            await db.execute(f'UPDATE sesiones SET {", ".join(fields)} WHERE id=?', values)
            await db.commit()

    return {'status': 'ok', 'sesionId': sesion_id}


@app.get('/api/sesion/{sesion_id}/status')
async def sesion_status(sesion_id: str):
    """Lightweight endpoint: member response status + processing progress.
    Safe to poll every few seconds from the wait panel."""
    async with aiosqlite.connect(DB_PATH) as db:
        db.row_factory = aiosqlite.Row
        cur = await db.execute('SELECT id,nombre,rol,email,respondio,email_sent FROM miembros WHERE sesion_id=? ORDER BY created_at', (sesion_id,))
        members = [dict(r) for r in await cur.fetchall()]
        cur2 = await db.execute('SELECT modelo,miembro_nombre,estado FROM procesamiento WHERE sesion_id=?', (sesion_id,))
        proc = [dict(r) for r in await cur2.fetchall()]
        cur3 = await db.execute('SELECT bala IS NOT NULL as has_bala FROM sesiones WHERE id=?', (sesion_id,))
        row = await cur3.fetchone()

    total      = len(members)
    respondidos = sum(1 for m in members if m['respondio'])
    proc_ok    = sum(1 for p in proc if p['estado'] == 'ok')
    return {
        'total':       total,
        'respondidos': respondidos,
        'allDone':     total > 0 and respondidos >= total,
        'members': [
            {
                'name':      m['nombre'],
                'role':      m['rol'],
                'email':     m['email'],
                'completed': bool(m['respondio']),
                'emailSent': bool(m['email_sent']),
            }
            for m in members
        ],
        'procesamiento': {
            'total': len(proc),
            'ok':    proc_ok,
            'done':  len(proc) > 0 and proc_ok >= len(proc),
        },
        'hasBrainstorm': bool(row and row['has_bala']),
    }


@app.post('/api/sesion/{sesion_id}/brainstorm')
async def save_brainstorm(sesion_id: str, req: Request):
    """Save or overwrite the brainstorming text (bala) for a session."""
    data = await req.json()
    texto = (data.get('texto') or data.get('text') or '').strip()
    if not texto:
        raise HTTPException(400, 'Texto vacÃ­o')
    async with aiosqlite.connect(DB_PATH) as db:
        cur = await db.execute('SELECT id FROM sesiones WHERE id=?', (sesion_id,))
        if not await cur.fetchone():
            raise HTTPException(404, 'SesiÃ³n no encontrada')
        await db.execute('UPDATE sesiones SET bala=? WHERE id=?', (texto, sesion_id))
        await db.commit()
    return {'status': 'ok', 'chars': len(texto)}


@app.post('/api/sesion/{sesion_id}/reenviar-emails')
async def reenviar_emails(sesion_id: str, req: Request):
    """Re-send invitation emails to members who haven't responded yet."""
    data     = await req.json()
    pub_url  = (data.get('publicUrl') or '').strip().rstrip('/')
    keys     = load_keys()
    smtp_cfg = {'user': keys.get('gmail_user',''), 'password': keys.get('gmail_password','')}

    if not smtp_cfg['user'] or not smtp_cfg['password']:
        raise HTTPException(400, 'Gmail no configurado en keys.json')

    async with aiosqlite.connect(DB_PATH) as db:
        db.row_factory = aiosqlite.Row
        cur = await db.execute('SELECT * FROM sesiones WHERE id=?', (sesion_id,))
        sesion = await cur.fetchone()
        if not sesion:
            raise HTTPException(404, 'SesiÃ³n no encontrada')
        sesion = dict(sesion)
        cur2 = await db.execute(
            'SELECT * FROM miembros WHERE sesion_id=? AND respondio=0 ORDER BY created_at', (sesion_id,))
        pendientes = [dict(r) for r in await cur2.fetchall()]

    ip        = get_local_ip()
    base_url  = pub_url if pub_url else f'http://{ip}:{PORT}'
    enviados, errores = 0, []

    for m in pendientes:
        if not m.get('email'):
            continue
        link = f'{base_url}/responder/{m["token"]}'
        try:
            send_email(m['email'], m['nombre'], link, sesion['problema'],
                       sesion.get('facilitador','REVOLVER'), smtp_cfg)
            enviados += 1
            async with aiosqlite.connect(DB_PATH) as db:
                await db.execute('UPDATE miembros SET email_sent=1, link=? WHERE id=?', (link, m['id']))
                await db.commit()
        except Exception as e:
            errores.append(f'{m["nombre"]}: {str(e)[:120]}')

    return {'status': 'ok', 'enviados': enviados, 'pendientes': len(pendientes), 'errores': errores}


# ââ Historial âââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââ

@app.get('/api/historial')
async def historial(authorization: Optional[str] = Header(None)):
    user = await _get_current_user(authorization)
    uid  = user['id'] if user else None
    async with aiosqlite.connect(DB_PATH) as db:
        db.row_factory = aiosqlite.Row
        if uid:
            cur = await db.execute(
                'SELECT id, empresa, problema, facilitador, estado, usuario_id, created_at, bala FROM sesiones WHERE usuario_id=? ORDER BY created_at DESC LIMIT 100',
                (uid,))
        else:
            cur = await db.execute(
                'SELECT id, empresa, problema, facilitador, estado, usuario_id, created_at, bala FROM sesiones ORDER BY created_at DESC LIMIT 100')
        sesiones = [dict(r) for r in await cur.fetchall()]
        # Enrich with member stats
        for s in sesiones:
            cur2 = await db.execute(
                'SELECT COUNT(*) as total, SUM(respondio) as respondidos FROM miembros WHERE sesion_id=?', (s['id'],))
            stats = dict(await cur2.fetchone())
            s['total_miembros']    = stats['total'] or 0
            s['respondidos']       = int(stats['respondidos'] or 0)
            s['tiene_brainstorm']  = bool(s.get('bala'))
            cur3 = await db.execute(
                'SELECT COUNT(*) as total, SUM(CASE WHEN estado="ok" THEN 1 ELSE 0 END) as ok FROM procesamiento WHERE sesion_id=?', (s['id'],))
            proc = dict(await cur3.fetchone())
            s['total_analisis']    = proc['total'] or 0
            s['analisis_ok']       = int(proc['ok'] or 0)
            s['problema_corto']    = (s['problema'] or '')[:120]
            del s['bala']  # don't send full brainstorm in list
    return {'sesiones': sesiones, 'total': len(sesiones)}


@app.delete('/api/sesion/{sesion_id}')
async def delete_sesion(sesion_id: str):
    async with aiosqlite.connect(DB_PATH) as db:
        await db.execute('DELETE FROM procesamiento WHERE sesion_id=?', (sesion_id,))
        await db.execute('DELETE FROM miembros WHERE sesion_id=?', (sesion_id,))
        await db.execute('DELETE FROM sesiones WHERE id=?', (sesion_id,))
        await db.commit()
    return {'status': 'ok'}


# ââ Responder (member submits) ââââââââââââââââââââââââââââââââââââââââââââââââ

@app.get('/api/sesion-by-token/{token}')
async def sesion_by_token(token: str):
    async with aiosqlite.connect(DB_PATH) as db:
        db.row_factory = aiosqlite.Row
        cur = await db.execute('SELECT * FROM miembros WHERE token=?', (token,))
        m = await cur.fetchone()
        if not m:
            raise HTTPException(404, 'Token no vÃ¡lido')
        m = dict(m)
        cur2 = await db.execute('SELECT * FROM sesiones WHERE id=?', (m['sesion_id'],))
        s = dict(await cur2.fetchone())
    dna = json.loads(s.get('dna', '{}'))
    return {
        'member': {'name': m['nombre'], 'role': m['rol']},
        'problem': s['problema'],
        'context': s.get('contexto', ''),
        'constraints': s.get('restricciones', ''),
        'success': s.get('criterios', ''),
        'empresa': dna.get('name', s.get('empresa', '')),
        'facilitador': s.get('facilitador', ''),
        'hasResponded': bool(m['respondio']),
    }


@app.get('/api/sesion-por-token/{token}')
async def sesion_por_token(token: str):
    """Endpoint used by responder.html â returns data in the format it expects."""
    async with aiosqlite.connect(DB_PATH) as db:
        db.row_factory = aiosqlite.Row
        cur = await db.execute('SELECT * FROM miembros WHERE token=?', (token,))
        m = await cur.fetchone()
        if not m:
            raise HTTPException(404, 'Token no vÃ¡lido')
        m = dict(m)
        cur2 = await db.execute('SELECT * FROM sesiones WHERE id=?', (m['sesion_id'],))
        s = dict(await cur2.fetchone())
    dna = json.loads(s.get('dna', '{}'))
    return {
        'member': {
            'name':      m['nombre'],
            'role':      m['rol'],
            'completed': bool(m['respondio']),
        },
        'sesion': {
            'problem':    s['problema'],
            'context':    s.get('contexto', ''),
            'constraints': s.get('restricciones', ''),
            'success':    s.get('criterios', ''),
            'brainstorm': s.get('bala') or '',
            'empresa':    dna.get('name', s.get('empresa', '')),
            'facilitador': s.get('facilitador', ''),
        },
    }


@app.post('/api/responder/{token}')
async def submit_respuesta(token: str, req: Request):
    data = await req.json()
    # Accept both 'response' (legacy) and 'input' (used by responder.html)
    resp = (data.get('input') or data.get('response') or '').strip()
    if not resp:
        raise HTTPException(400, 'Respuesta vacÃ­a')
    async with aiosqlite.connect(DB_PATH) as db:
        cur = await db.execute('SELECT id FROM miembros WHERE token=?', (token,))
        m = await cur.fetchone()
        if not m:
            raise HTTPException(404, 'Token no vÃ¡lido')
        await db.execute('UPDATE miembros SET respondio=1, respuesta=? WHERE token=?',
                         (resp, token))
        await db.commit()
    return {'status': 'ok', 'message': 'Â¡Gracias! Tu perspectiva fue registrada.'}


# ââ WebSocket: Procesamiento en tiempo real âââââââââââââââââââââââââââââââââââ

@app.websocket('/ws/procesar/{sesion_id}')
async def ws_procesar(websocket: WebSocket, sesion_id: str):
    await ws_mgr.connect(sesion_id, websocket)
    try:
        # Start processing in background
        asyncio.create_task(process_all(sesion_id))
        # Keep connection alive until client disconnects
        while True:
            await websocket.receive_text()
    except WebSocketDisconnect:
        ws_mgr.disconnect(sesion_id, websocket)
        log.info(f'WS desconectado sesion={sesion_id}')


# ââ WebSocket: La Bala en streaming ââââââââââââââââââââââââââââââââââââââââââ

@app.websocket('/ws/bala/{sesion_id}')
async def ws_bala(websocket: WebSocket, sesion_id: str):
    await ws_mgr.connect(sesion_id, websocket)
    try:
        sesion = await db_get_sesion(sesion_id)
        if not sesion:
            await websocket.send_json({'type': 'bala_error', 'msg': 'SesiÃ³n no encontrada'})
            return
        asyncio.create_task(synthesize_stream(sesion, sesion_id))
        while True:
            await websocket.receive_text()
    except WebSocketDisconnect:
        ws_mgr.disconnect(sesion_id, websocket)


# ââ Test endpoints ââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââ

@app.post('/api/test-models')
async def test_models():
    keys = load_keys()
    mini = 'Di solo "ok".'
    results = {}
    async with httpx.AsyncClient(timeout=30) as client:
        tasks = {name: fn(mini, client, keys) for name, fn in AI_FUNCTIONS.items()}
        for name, coro in tasks.items():
            try:
                r = await coro
                results[name] = {'status': r.get('status'), 'detail': r.get('result', '')[:120]}
            except Exception as e:
                results[name] = {'status': 'exception', 'detail': str(e)[:120]}
    return {'results': results}


@app.post('/api/extract-text')
async def extract_text(file: UploadFile = File(...)):
    """Extract plain text from PDF, DOCX, XLSX, TXT files for AI context."""
    ext = (file.filename or '').rsplit('.', 1)[-1].lower()
    data = await file.read()
    try:
        if ext in ('txt', 'md', 'csv'):
            text = data.decode('utf-8', errors='replace')

        elif ext == 'xlsx' or ext == 'xls':
            import openpyxl, io
            wb = openpyxl.load_workbook(io.BytesIO(data), data_only=True)
            parts = []
            for sheet in wb.worksheets:
                parts.append(f'=== Hoja: {sheet.title} ===')
                for row in sheet.iter_rows(values_only=True):
                    cells = [str(c) if c is not None else '' for c in row]
                    if any(c.strip() for c in cells):
                        parts.append('\t'.join(cells))
            text = '\n'.join(parts)

        elif ext in ('docx', 'doc'):
            import docx as _docx, io
            doc = _docx.Document(io.BytesIO(data))
            text = '\n'.join(p.text for p in doc.paragraphs if p.text.strip())

        elif ext == 'pdf':
            import fitz, io
            pdf = fitz.open(stream=data, filetype='pdf')
            text = '\n'.join(page.get_text() for page in pdf)

        else:
            raise HTTPException(400, f'Formato no soportado: {ext}')

        return {'text': text[:50000], 'chars': len(text), 'filename': file.filename}

    except HTTPException:
        raise
    except Exception as e:
        log.error(f'extract-text error: {e}')
        raise HTTPException(500, f'Error al procesar archivo: {e}')


@app.post('/api/test-email')
async def test_email(req: Request):
    data = await req.json()
    keys = load_keys()
    smtp = {'user': keys.get('gmail_user', ''), 'password': keys.get('gmail_password', '')}
    if not smtp['user'] or not smtp['password']:
        return {'status': 'error', 'error': 'Gmail no configurado'}
    to = (data.get('email') or smtp['user'])
    ip = get_local_ip()
    try:
        send_email(to, 'Test Usuario', f'http://{ip}:{PORT}/responder/TEST',
                   'Email de prueba REVOLVER.', 'REVOLVER Admin', smtp)
        return {'status': 'ok', 'message': f'Email enviado a {to}'}
    except Exception as e:
        return {'status': 'error', 'error': str(e)[:300]}


# ââ Save DNA (persiste en keys-side; guardado en sesiÃ³n vÃ­a crear-sesion) âââââ

@app.post('/api/save-dna')
async def save_dna(req: Request):
    """Guarda el DNA globalmente para que persista entre sesiones."""
    data = await req.json()
    # Guardar en un archivo dna.json para persistencia global
    dna_path = os.path.join(BASE_DIR, 'dna.json')
    with open(dna_path, 'w', encoding='utf-8') as f:
        json.dump(data, f, ensure_ascii=False, indent=2)
    return {'status': 'ok'}


@app.get('/api/load-dna')
async def load_dna():
    dna_path = os.path.join(BASE_DIR, 'dna.json')
    if os.path.exists(dna_path):
        with open(dna_path, encoding='utf-8') as f:
            return json.load(f)
    return {}


# ââ Export DOCX âââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââ

@app.post('/api/export-docx')
async def export_docx(req: Request):
    """Genera un Word .docx ejecutivo del Advisory Brief."""
    try:
        from docx import Document
        from docx.shared import Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        import re, io
    except ImportError:
        return JSONResponse({'error': 'python-docx no instalado. Ejecuta: pip install python-docx'}, 400)

    data     = await req.json()
    md_text  = data.get('text', '')
    empresa  = data.get('empresa', 'Empresa')
    problema = data.get('problema', '')
    fecha    = datetime.utcnow().strftime('%d/%m/%Y')

    # ââ Remove DATOS_VISUALES block ââââââââââââââââââââââââââââââââââââââââââ
    idx = md_text.find('## DATOS_VISUALES')
    if idx != -1:
        md_text = md_text[:idx].rstrip()

    # ââ Build document âââââââââââââââââââââââââââââââââââââââââââââââââââââââ
    doc = Document()

    # Page margins
    for sec in doc.sections:
        sec.top_margin    = Cm(2.5)
        sec.bottom_margin = Cm(2.5)
        sec.left_margin   = Cm(2.8)
        sec.right_margin  = Cm(2.8)

    # ââ Styles âââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââ
    def style_normal(para, size=11, bold=False, italic=False, color=None, space_after=6):
        para.paragraph_format.space_after  = Pt(space_after)
        para.paragraph_format.space_before = Pt(0)
        for run in para.runs:
            run.font.size   = Pt(size)
            run.font.bold   = bold
            run.font.italic = italic
            if color:
                run.font.color.rgb = RGBColor(*color)

    def add_para(text, style='Normal', align=None):
        p = doc.add_paragraph(text, style=style)
        if align:
            p.alignment = align
        return p

    def add_run_inline(para, text):
        """Add inline markdown (bold/italic) to an existing paragraph."""
        parts = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)', text)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                r = para.add_run(part[2:-2]); r.bold = True
            elif part.startswith('*') and part.endswith('*'):
                r = para.add_run(part[1:-1]); r.italic = True
            elif part.startswith('`') and part.endswith('`'):
                r = para.add_run(part[1:-1])
                r.font.name = 'Courier New'; r.font.size = Pt(9)
            else:
                para.add_run(part)

    # ââ Cover page âââââââââââââââââââââââââââââââââââââââââââââââââââââââââââ
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(80)
    p.paragraph_format.space_after  = Pt(8)
    r = p.add_run('REVOLVER')
    r.font.name  = 'Courier New'
    r.font.size  = Pt(9)
    r.font.color.rgb = RGBColor(200, 16, 46)
    r.font.bold  = True
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(4)
    r2 = p2.add_run('Sistema Multi-IA de ConsultorÃ­a EstratÃ©gica')
    r2.font.name = 'Courier New'; r2.font.size = Pt(8); r2.font.color.rgb = RGBColor(100,100,100)

    p_name = doc.add_paragraph()
    p_name.paragraph_format.space_before = Pt(24)
    p_name.paragraph_format.space_after  = Pt(16)
    r_name = p_name.add_run(empresa.upper())
    r_name.font.size = Pt(32); r_name.bold = True
    r_name.font.color.rgb = RGBColor(14, 15, 18)

    if problema:
        p_prob = doc.add_paragraph()
        p_prob.paragraph_format.space_after = Pt(12)
        r_prob = p_prob.add_run(problema[:200])
        r_prob.font.size = Pt(12); r_prob.font.color.rgb = RGBColor(80,80,80)

    p_date = doc.add_paragraph()
    p_date.paragraph_format.space_before = Pt(32)
    r_date = p_date.add_run(f'Advisory Brief Â· {fecha} Â· Confidencial')
    r_date.font.name = 'Courier New'; r_date.font.size = Pt(8)
    r_date.font.color.rgb = RGBColor(150,120,50)

    doc.add_page_break()

    # ââ Parse markdown and build document ââââââââââââââââââââââââââââââââââââ
    lines      = md_text.split('\n')
    in_table   = False
    table_rows = []

    def flush_table():
        nonlocal in_table, table_rows
        if not table_rows: return
        # Filter separator rows
        data_rows = [r for r in table_rows if not all(c.strip().startswith('-') for c in r)]
        if not data_rows:
            in_table = False; table_rows = []; return
        cols = max(len(r) for r in data_rows)
        t = doc.add_table(rows=len(data_rows), cols=cols)
        t.style = 'Table Grid'
        for ri, row in enumerate(data_rows):
            for ci, cell_text in enumerate(row[:cols]):
                cell = t.cell(ri, ci)
                cell.text = cell_text.strip()
                p = cell.paragraphs[0]
                p.paragraph_format.space_after  = Pt(2)
                p.paragraph_format.space_before = Pt(2)
                for run in p.runs:
                    run.font.size = Pt(9)
                    if ri == 0:
                        run.bold = True
                        run.font.color.rgb = RGBColor(255,255,255)
                if ri == 0:
                    tc = cell._tc
                    tcPr = tc.get_or_add_tcPr()
                    shd = OxmlElement('w:shd')
                    shd.set(qn('w:val'), 'clear')
                    shd.set(qn('w:color'), 'auto')
                    shd.set(qn('w:fill'), '0E0F12')
                    tcPr.append(shd)
        doc.add_paragraph().paragraph_format.space_after = Pt(6)
        in_table = False; table_rows = []

    in_list = False
    for raw in lines:
        line = raw.rstrip()

        # Table
        if line.startswith('|'):
            if not in_table:
                flush_table()
                in_table = True; table_rows = []
            cells = line.split('|')[1:-1]
            table_rows.append(cells)
            continue
        elif in_table:
            flush_table()

        # Lists
        if re.match(r'^[-*â¢]\s', line):
            in_list = True
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_after = Pt(3)
            add_run_inline(p, re.sub(r'^[-*â¢]\s+', '', line))
            continue
        elif re.match(r'^\d+\.\s', line):
            in_list = True
            p = doc.add_paragraph(style='List Number')
            p.paragraph_format.space_after = Pt(3)
            add_run_inline(p, re.sub(r'^\d+\.\s+', '', line))
            continue
        else:
            in_list = False

        # Headings
        if re.match(r'^# ', line):
            p = doc.add_heading(line[2:], level=1)
            p.paragraph_format.space_before = Pt(20)
            p.paragraph_format.space_after  = Pt(8)
            for run in p.runs:
                run.font.size  = Pt(22); run.font.color.rgb = RGBColor(14,15,18)
        elif re.match(r'^## ', line):
            p = doc.add_heading(line[3:], level=2)
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after  = Pt(6)
            for run in p.runs:
                run.font.size  = Pt(13); run.font.color.rgb = RGBColor(154,123,58)
        elif re.match(r'^### ', line):
            p = doc.add_heading(line[4:], level=3)
            p.paragraph_format.space_after = Pt(4)
            for run in p.runs:
                run.font.size  = Pt(11); run.font.color.rgb = RGBColor(14,15,18)
        elif re.match(r'^---+$', line):
            p = doc.add_paragraph()
            p.paragraph_format.space_after  = Pt(8)
            p.paragraph_format.space_before = Pt(8)
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), '4')
            bottom.set(qn('w:space'), '1'); bottom.set(qn('w:color'), 'D8D4CC')
            pBdr.append(bottom)
            pPr.append(pBdr)
        elif not line.strip():
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            add_run_inline(p, line)
            for run in p.runs:
                run.font.size = Pt(10.5)

    flush_table()

    # ââ Footer on last page âââââââââââââââââââââââââââââââââââââââââââââââââââ
    doc.add_paragraph()
    p_footer = doc.add_paragraph()
    p_footer.paragraph_format.space_before = Pt(24)
    r_f = p_footer.add_run(f'Generado por REVOLVER Â· Sistema Multi-IA de ConsultorÃ­a Â· {fecha}')
    r_f.font.name = 'Courier New'; r_f.font.size = Pt(7.5)
    r_f.font.color.rgb = RGBColor(150,148,144)

    # ââ Return as file ââââââââââââââââââââââââââââââââââââââââââââââââââââââââ
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    filename = f"REVOLVER_Advisory_{empresa.replace(' ','_')}_{datetime.utcnow().strftime('%Y%m%d')}.docx"
    from fastapi.responses import StreamingResponse
    return StreamingResponse(
        buf,
        media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        headers={'Content-Disposition': f'attachment; filename="{filename}"'}
    )


# ââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââ
# AUTH â usuarios/empresas
# ââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââ

def _hash_password(password: str, salt: str = None):
    if salt is None:
        salt = secrets.token_hex(16)
    key = hashlib.pbkdf2_hmac('sha256', password.encode('utf-8'), salt.encode('utf-8'), 100_000)
    return key.hex(), salt


def _verify_password(password: str, stored_hash: str, salt: str) -> bool:
    computed, _ = _hash_password(password, salt)
    return hmac.compare_digest(computed, stored_hash)


async def _create_auth_token(usuario_id: str) -> str:
    token = secrets.token_hex(32)
    now   = datetime.utcnow()
    exp   = (now + timedelta(days=30)).isoformat()
    async with aiosqlite.connect(DB_PATH) as db:
        await db.execute(
            'INSERT INTO tokens_auth (token, usuario_id, expires_at, created_at) VALUES (?,?,?,?)',
            (token, usuario_id, exp, now.isoformat()))
        await db.commit()
    return token


async def _get_current_user(authorization: Optional[str] = Header(None)) -> Optional[dict]:
    if not authorization:
        return None
    raw = authorization.strip()
    token = raw[7:] if raw.lower().startswith('bearer ') else raw
    now = datetime.utcnow().isoformat()
    async with aiosqlite.connect(DB_PATH) as db:
        db.row_factory = aiosqlite.Row
        cur = await db.execute(
            '''SELECT u.* FROM usuarios u
               JOIN tokens_auth t ON u.id = t.usuario_id
               WHERE t.token=? AND t.expires_at > ?''',
            (token, now))
        row = await cur.fetchone()
    return dict(row) if row else None


async def _require_user(authorization: Optional[str] = Header(None)) -> dict:
    user = await _get_current_user(authorization)
    if not user:
        raise HTTPException(401, 'No autorizado â inicia sesiÃ³n primero')
    return user


# ââ Register ââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââ

@app.post('/api/auth/registrar')
async def auth_registrar(req: Request):
    data     = await req.json()
    nombre   = (data.get('nombre') or '').strip()
    empresa  = (data.get('empresa') or '').strip()
    cargo    = (data.get('cargo') or '').strip()
    email    = (data.get('email') or '').strip().lower()
    password = (data.get('password') or '').strip()
    industria = (data.get('industria') or '').strip()
    website  = (data.get('website') or '').strip()

    if not all([nombre, empresa, email, password]):
        raise HTTPException(400, 'Nombre, empresa, email y contraseÃ±a son requeridos')
    if len(password) < 6:
        raise HTTPException(400, 'La contraseÃ±a debe tener al menos 6 caracteres')

    async with aiosqlite.connect(DB_PATH) as db:
        db.row_factory = aiosqlite.Row
        cur = await db.execute('SELECT id FROM usuarios WHERE email=?', (email,))
        if await cur.fetchone():
            raise HTTPException(409, 'Ya existe una cuenta con ese email')

        uid   = secrets.token_hex(10)
        now   = datetime.utcnow().isoformat()
        phash, salt = _hash_password(password)
        await db.execute(
            '''INSERT INTO usuarios
               (id,nombre,empresa,cargo,email,password_hash,salt,industria,website,created_at)
               VALUES (?,?,?,?,?,?,?,?,?,?)''',
            (uid, nombre, empresa, cargo, email, phash, salt, industria, website, now))
        await db.commit()

    token = await _create_auth_token(uid)
    log.info(f'Usuario registrado: {email} empresa={empresa}')
    return {
        'status': 'ok',
        'token': token,
        'user': {'id': uid, 'nombre': nombre, 'empresa': empresa,
                 'cargo': cargo, 'email': email, 'industria': industria}
    }


@app.post('/api/auth/login')
async def auth_login(req: Request):
    data     = await req.json()
    email    = (data.get('email') or '').strip().lower()
    password = (data.get('password') or '').strip()

    if not email or not password:
        raise HTTPException(400, 'Email y contraseÃ±a requeridos')

    async with aiosqlite.connect(DB_PATH) as db:
        db.row_factory = aiosqlite.Row
        cur = await db.execute('SELECT * FROM usuarios WHERE email=?', (email,))
        row = await cur.fetchone()

    if not row:
        raise HTTPException(401, 'Email o contraseÃ±a incorrectos')
    user = dict(row)
    if not _verify_password(password, user['password_hash'], user['salt']):
        raise HTTPException(401, 'Email o contraseÃ±a incorrectos')

    token = await _create_auth_token(user['id'])
    log.info(f'Login: {email}')
    return {
        'status': 'ok',
        'token': token,
        'user': {
            'id': user['id'], 'nombre': user['nombre'], 'empresa': user['empresa'],
            'cargo': user['cargo'], 'email': user['email'],
            'industria': user['industria'], 'website': user['website'],
            'plan': user['plan'], 'logo_url': user['logo_url'],
        }
    }


@app.get('/api/auth/me')
async def auth_me(user: dict = Depends(_require_user)):
    return {
        'id': user['id'], 'nombre': user['nombre'], 'empresa': user['empresa'],
        'cargo': user['cargo'], 'email': user['email'],
        'industria': user['industria'], 'website': user['website'],
        'plan': user['plan'], 'logo_url': user['logo_url'],
    }


@app.put('/api/auth/perfil')
async def auth_update_perfil(req: Request, user: dict = Depends(_require_user)):
    data = await req.json()
    allowed = {'nombre': 'nombre', 'empresa': 'empresa', 'cargo': 'cargo',
               'industria': 'industria', 'website': 'website', 'logo_url': 'logo_url'}
    fields, values = [], []
    for k, col in allowed.items():
        if k in data:
            fields.append(f'{col}=?')
            values.append(str(data[k])[:300])
    if 'password' in data and len(data['password']) >= 6:
        phash, salt = _hash_password(data['password'])
        fields += ['password_hash=?', 'salt=?']
        values += [phash, salt]
    if fields:
        values.append(user['id'])
        async with aiosqlite.connect(DB_PATH) as db:
            await db.execute(f'UPDATE usuarios SET {", ".join(fields)} WHERE id=?', values)
            await db.commit()
    return {'status': 'ok'}


@app.post('/api/auth/logout')
async def auth_logout(authorization: Optional[str] = Header(None)):
    if authorization:
        raw = authorization.strip()
        token = raw[7:] if raw.lower().startswith('bearer ') else raw
        async with aiosqlite.connect(DB_PATH) as db:
            await db.execute('DELETE FROM tokens_auth WHERE token=?', (token,))
            await db.commit()
    return {'status': 'ok'}


# ââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââ
# SESIONES AVANZADAS â clonar, estado, miembros, ngrok URL
# ââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââ

@app.post('/api/sesion/{sesion_id}/clonar')
async def clonar_sesion(sesion_id: str, req: Request):
    """Duplica una sesiÃ³n (problema, DNA, miembros) creando una nueva con estado borrador."""
    data = await req.json()
    nuevo_problema = (data.get('problema') or '').strip()

    s = await db_get_sesion(sesion_id)
    if not s:
        raise HTTPException(404, 'SesiÃ³n origen no encontrada')

    new_id = secrets.token_hex(8)
    now    = datetime.utcnow().isoformat()
    ip     = get_local_ip()
    public_url = (data.get('publicUrl') or '').strip().rstrip('/')
    base_url   = public_url if public_url else f'http://{ip}:{PORT}'

    usuario_id = data.get('usuarioId') or s.get('usuario_id')
    empresa    = s['dna'].get('name') or s['empresa']

    async with aiosqlite.connect(DB_PATH) as db:
        await db.execute(
            '''INSERT INTO sesiones
               (id,empresa,problema,contexto,impacto,restricciones,criterios,
                facilitador,dna,estado,usuario_id,created_at)
               VALUES (?,?,?,?,?,?,?,?,?,?,?,?)''',
            (new_id, empresa,
             nuevo_problema or s['problema'],
             s.get('contexto',''), s.get('impacto',''),
             s.get('restricciones',''), s.get('criterios',''),
             s.get('facilitador','El facilitador'),
             json.dumps(s['dna']), 'borrador', usuario_id, now))
        await db.commit()

        # Clone members without their responses
        for m in s['members']:
            token = secrets.token_hex(16)
            link  = f'{base_url}/responder/{token}'
            await db.execute(
                '''INSERT INTO miembros
                   (id,sesion_id,nombre,rol,email,token,link,created_at)
                   VALUES (?,?,?,?,?,?,?,?)''',
                (secrets.token_hex(6), new_id, m['nombre'], m['rol'],
                 m.get('email',''), token, link, now))
        await db.commit()

    new_s = await db_get_sesion(new_id)
    log.info(f'SesiÃ³n clonada: {sesion_id} â {new_id}')
    return {'status': 'ok', 'sesionId': new_id, 'sesion': new_s}


@app.put('/api/sesion/{sesion_id}/estado')
async def update_estado(sesion_id: str, req: Request):
    """Cambia el estado de la sesiÃ³n: borrador | activa | completada | archivada"""
    data   = await req.json()
    estado = (data.get('estado') or '').strip()
    valid  = {'borrador', 'activa', 'completada', 'archivada'}
    if estado not in valid:
        raise HTTPException(400, f'Estado invÃ¡lido. VÃ¡lidos: {valid}')
    async with aiosqlite.connect(DB_PATH) as db:
        cur = await db.execute('SELECT id FROM sesiones WHERE id=?', (sesion_id,))
        if not await cur.fetchone():
            raise HTTPException(404, 'SesiÃ³n no encontrada')
        await db.execute('UPDATE sesiones SET estado=? WHERE id=?', (estado, sesion_id))
        await db.commit()
    return {'status': 'ok', 'estado': estado}


@app.post('/api/sesion/{sesion_id}/actualizar-url')
async def actualizar_url(sesion_id: str, req: Request):
    """Actualiza la base URL de todos los links de miembros (Ãºtil cuando cambia ngrok)."""
    data       = await req.json()
    public_url = (data.get('publicUrl') or '').strip().rstrip('/')
    reenviar   = bool(data.get('reenviar', False))

    if not public_url:
        raise HTTPException(400, 'publicUrl es requerido')

    keys     = load_keys()
    smtp_cfg = {'user': keys.get('gmail_user',''), 'password': keys.get('gmail_password','')}

    async with aiosqlite.connect(DB_PATH) as db:
        db.row_factory = aiosqlite.Row
        cur = await db.execute('SELECT * FROM sesiones WHERE id=?', (sesion_id,))
        sesion = await cur.fetchone()
        if not sesion:
            raise HTTPException(404, 'SesiÃ³n no encontrada')
        sesion = dict(sesion)

        cur2 = await db.execute('SELECT * FROM miembros WHERE sesion_id=?', (sesion_id,))
        miembros = [dict(r) for r in await cur2.fetchall()]

    actualizados, enviados, errores = 0, 0, []
    for m in miembros:
        new_link = f'{public_url}/responder/{m["token"]}'
        async with aiosqlite.connect(DB_PATH) as db:
            await db.execute('UPDATE miembros SET link=? WHERE id=?', (new_link, m['id']))
            await db.commit()
        actualizados += 1

        if reenviar and not m.get('respondio') and m.get('email') and smtp_cfg['user']:
            try:
                send_email(m['email'], m['nombre'], new_link, sesion['problema'],
                           sesion.get('facilitador','REVOLVER'), smtp_cfg)
                enviados += 1
                async with aiosqlite.connect(DB_PATH) as db:
                    await db.execute('UPDATE miembros SET email_sent=1 WHERE id=?', (m['id'],))
                    await db.commit()
            except Exception as e:
                errores.append(f'{m["nombre"]}: {str(e)[:100]}')

    log.info(f'URL actualizada sesion={sesion_id} nuevos_links={actualizados}')
    return {'status': 'ok', 'actualizados': actualizados, 'emailsEnviados': enviados, 'errores': errores}


@app.post('/api/sesion/{sesion_id}/miembros')
async def agregar_miembro(sesion_id: str, req: Request):
    """Agrega un miembro nuevo a una sesiÃ³n existente."""
    data       = await req.json()
    nombre     = (data.get('name') or data.get('nombre') or '').strip()
    rol        = (data.get('role') or data.get('rol') or 'General').strip()
    email      = (data.get('email') or '').strip()
    public_url = (data.get('publicUrl') or '').strip().rstrip('/')

    if not nombre:
        raise HTTPException(400, 'Nombre es requerido')

    ip       = get_local_ip()
    base_url = public_url if public_url else f'http://{ip}:{PORT}'
    keys     = load_keys()
    smtp_cfg = {'user': keys.get('gmail_user',''), 'password': keys.get('gmail_password','')}

    async with aiosqlite.connect(DB_PATH) as db:
        db.row_factory = aiosqlite.Row
        cur = await db.execute('SELECT * FROM sesiones WHERE id=?', (sesion_id,))
        sesion = await cur.fetchone()
        if not sesion:
            raise HTTPException(404, 'SesiÃ³n no encontrada')
        sesion = dict(sesion)

    now   = datetime.utcnow().isoformat()
    token = secrets.token_hex(16)
    link  = f'{base_url}/responder/{token}'
    mid   = secrets.token_hex(6)

    async with aiosqlite.connect(DB_PATH) as db:
        await db.execute(
            'INSERT INTO miembros (id,sesion_id,nombre,rol,email,token,link,created_at) VALUES (?,?,?,?,?,?,?,?)',
            (mid, sesion_id, nombre, rol, email, token, link, now))
        await db.commit()

    email_sent = False
    if email and smtp_cfg['user'] and smtp_cfg['password']:
        try:
            send_email(email, nombre, link, sesion['problema'],
                       sesion.get('facilitador','REVOLVER'), smtp_cfg)
            email_sent = True
            async with aiosqlite.connect(DB_PATH) as db:
                await db.execute('UPDATE miembros SET email_sent=1 WHERE id=?', (mid,))
                await db.commit()
        except Exception as e:
            log.warning(f'Email agregar_miembro fail: {e}')

    return {
        'status': 'ok',
        'member': {
            'id': mid, 'name': nombre, 'role': rol, 'email': email,
            'token': token, 'link': link, 'emailSent': email_sent,
            'completed': False,
        }
    }


@app.delete('/api/sesion/{sesion_id}/miembros/{member_id}')
async def eliminar_miembro(sesion_id: str, member_id: str):
    """Elimina un miembro de una sesiÃ³n (solo si no ha respondido)."""
    async with aiosqlite.connect(DB_PATH) as db:
        db.row_factory = aiosqlite.Row
        cur = await db.execute(
            'SELECT id, respondio FROM miembros WHERE id=? AND sesion_id=?', (member_id, sesion_id))
        m = await cur.fetchone()
        if not m:
            raise HTTPException(404, 'Miembro no encontrado')
        if dict(m)['respondio']:
            raise HTTPException(400, 'No se puede eliminar un miembro que ya respondiÃ³')
        await db.execute('DELETE FROM miembros WHERE id=?', (member_id,))
        await db.commit()
    return {'status': 'ok'}


# ââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââ
# TEMPLATES â problemas precargados por tipo
# ââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââ

# Default system templates seeded on first run
DEFAULT_TEMPLATES = [
    {
        'titulo': 'Crisis de ReputaciÃ³n',
        'tipo': 'crisis',
        'descripcion': 'DaÃ±o a imagen pÃºblica por incidente externo o comunicacional',
        'problema': 'La organizaciÃ³n enfrenta una crisis de reputaciÃ³n que amenaza su posiciÃ³n en el mercado y la confianza de sus stakeholders.',
        'contexto': 'La situaciÃ³n escalÃ³ rÃ¡pidamente en medios y redes sociales. El equipo directivo debe actuar con rapidez y coherencia.',
        'impacto': 'PÃ©rdida de clientes, caÃ­da en ventas, dificultad para retener talento y tensiÃ³n con reguladores.',
        'restricciones': 'No se pueden hacer declaraciones pÃºblicas sin validaciÃ³n legal. Plazo mÃ¡ximo de respuesta: 48 horas.',
        'criterios': 'RecuperaciÃ³n de NPS en 90 dÃ­as, cero nuevas menciones negativas virales, plan comunicacional ejecutado.',
    },
    {
        'titulo': 'DisrupciÃ³n TecnolÃ³gica',
        'tipo': 'innovacion',
        'descripcion': 'Un competidor o nueva tecnologÃ­a amenaza el modelo de negocio actual',
        'problema': 'Una tecnologÃ­a disruptiva estÃ¡ redefiniendo el sector. La organizaciÃ³n debe decidir si adopta, adapta o defiende su posiciÃ³n actual.',
        'contexto': 'Los competidores emergentes capturan cuota de mercado con soluciones mÃ¡s Ã¡giles y costos menores.',
        'impacto': 'PÃ©rdida progresiva de relevancia, presiÃ³n en mÃ¡rgenes y riesgo de obsolescencia en 24â36 meses.',
        'restricciones': 'Presupuesto de transformaciÃ³n limitado. No se pueden discontinuar lÃ­neas de negocio core sin plan de transiciÃ³n.',
        'criterios': 'Hoja de ruta tecnolÃ³gica aprobada en 60 dÃ­as, primer MVP en 180 dÃ­as, KPIs de adopciÃ³n definidos.',
    },
    {
        'titulo': 'ExpansiÃ³n a Nuevo Mercado',
        'tipo': 'estrategia',
        'descripcion': 'Ingreso a un mercado geogrÃ¡fico o segmento no atendido',
        'problema': 'La organizaciÃ³n evalÃºa expandirse a un nuevo mercado con alto potencial pero con riesgos regulatorios, culturales y operacionales significativos.',
        'contexto': 'El mercado objetivo muestra tasas de crecimiento superiores al mercado domÃ©stico pero requiere adaptaciÃ³n del modelo de negocio.',
        'impacto': 'Potencial de incrementar ingresos en 30â50% en 3 aÃ±os, con una inversiÃ³n inicial de alto riesgo.',
        'restricciones': 'Capacidad operacional actual al 85%. Equipo directivo sin experiencia en el mercado objetivo.',
        'criterios': 'DecisiÃ³n de go/no-go en 45 dÃ­as, plan de entrada validado, socios locales identificados.',
    },
    {
        'titulo': 'ReestructuraciÃ³n Organizacional',
        'tipo': 'operaciones',
        'descripcion': 'RediseÃ±o de estructura, procesos o cultura ante cambio de contexto',
        'problema': 'La estructura organizacional actual es un obstÃ¡culo para la velocidad de decisiÃ³n y la ejecuciÃ³n estratÃ©gica.',
        'contexto': 'El crecimiento acelerado creÃ³ silos funcionales, duplicidad de roles y procesos lentos que afectan la competitividad.',
        'impacto': 'Tiempo de respuesta al mercado 3x mÃ¡s lento que competidores. RotaciÃ³n del talento clave por encima del 20% anual.',
        'restricciones': 'No se pueden hacer reducciones de dotaciÃ³n. El cambio debe ser progresivo para no afectar la operaciÃ³n.',
        'criterios': 'Nueva estructura implementada en 120 dÃ­as, reducciÃ³n del 40% en ciclos de aprobaciÃ³n, NPS interno sobre 60.',
    },
    {
        'titulo': 'Problema Financiero CrÃ­tico',
        'tipo': 'finanzas',
        'descripcion': 'PresiÃ³n de liquidez, caÃ­da de mÃ¡rgenes o deuda insostenible',
        'problema': 'La organizaciÃ³n enfrenta una situaciÃ³n financiera crÃ­tica que requiere decisiones urgentes para asegurar la continuidad operacional.',
        'contexto': 'CombinaciÃ³n de factores externos (mercado) e internos (costos) generaron una presiÃ³n de caja no anticipada.',
        'impacto': 'Runway de caja proyectado en menos de 6 meses si no se toman acciones inmediatas.',
        'restricciones': 'No hay apetito de los accionistas para aportes de capital adicionales. Covenant bancario en riesgo.',
        'criterios': 'Plan de estabilizaciÃ³n financiera en 30 dÃ­as, runway extendido a 18 meses, covenant normalizado en 90 dÃ­as.',
    },
    {
        'titulo': 'VUCA â Incertidumbre EstratÃ©gica',
        'tipo': 'vuca',
        'descripcion': 'Entorno volÃ¡til, incierto, complejo y ambiguo',
        'problema': 'La organizaciÃ³n opera en un entorno VUCA donde las herramientas tradicionales de planificaciÃ³n resultan insuficientes.',
        'contexto': 'Alta volatilidad regulatoria, macroeconÃ³mica y competitiva simultÃ¡nea que invalida los supuestos del plan estratÃ©gico.',
        'impacto': 'Incapacidad de comprometer inversiones a mÃ¡s de 12 meses, parÃ¡lisis decisional en la alta direcciÃ³n.',
        'restricciones': 'No se puede esperar a tener certeza completa. Las decisiones deben tomarse con informaciÃ³n parcial.',
        'criterios': 'Marco de decisiÃ³n bajo incertidumbre adoptado, escenarios actualizados mensualmente, equipo alineado.',
    },
]

async def _seed_templates():
    async with aiosqlite.connect(DB_PATH) as db:
        cur = await db.execute('SELECT COUNT(*) as c FROM templates WHERE usuario_id IS NULL')
        row = await cur.fetchone()
        count = row[0] if row else 0
        if count == 0:
            now = datetime.utcnow().isoformat()
            for t in DEFAULT_TEMPLATES:
                tid = secrets.token_hex(8)
                await db.execute(
                    '''INSERT INTO templates
                       (id,titulo,tipo,descripcion,problema,contexto,impacto,
                        restricciones,criterios,usuario_id,created_at)
                       VALUES (?,?,?,?,?,?,?,?,?,NULL,?)''',
                    (tid, t['titulo'], t['tipo'], t.get('descripcion',''),
                     t['problema'], t['contexto'], t['impacto'],
                     t['restricciones'], t['criterios'], now))
            await db.commit()
            log.info(f'Templates sembrados: {len(DEFAULT_TEMPLATES)} plantillas')


@app.get('/api/templates')
async def get_templates(authorization: Optional[str] = Header(None)):
    user = await _get_current_user(authorization)
    uid  = user['id'] if user else None
    async with aiosqlite.connect(DB_PATH) as db:
        db.row_factory = aiosqlite.Row
        if uid:
            cur = await db.execute(
                'SELECT * FROM templates WHERE usuario_id IS NULL OR usuario_id=? ORDER BY tipo, titulo',
                (uid,))
        else:
            cur = await db.execute(
                'SELECT * FROM templates WHERE usuario_id IS NULL ORDER BY tipo, titulo')
        rows = [dict(r) for r in await cur.fetchall()]
    return {'templates': rows, 'total': len(rows)}


@app.post('/api/templates')
async def create_template(req: Request, user: dict = Depends(_require_user)):
    data = await req.json()
    titulo = (data.get('titulo') or '').strip()
    if not titulo:
        raise HTTPException(400, 'TÃ­tulo requerido')
    tid = secrets.token_hex(8)
    now = datetime.utcnow().isoformat()
    async with aiosqlite.connect(DB_PATH) as db:
        await db.execute(
            '''INSERT INTO templates
               (id,titulo,tipo,descripcion,problema,contexto,impacto,
                restricciones,criterios,usuario_id,created_at)
               VALUES (?,?,?,?,?,?,?,?,?,?,?)''',
            (tid, titulo, data.get('tipo','general'), data.get('descripcion',''),
             data.get('problema',''), data.get('contexto',''), data.get('impacto',''),
             data.get('restricciones',''), data.get('criterios',''), user['id'], now))
        await db.commit()
    return {'status': 'ok', 'id': tid}


@app.delete('/api/templates/{template_id}')
async def delete_template(template_id: str, user: dict = Depends(_require_user)):
    async with aiosqlite.connect(DB_PATH) as db:
        cur = await db.execute(
            'SELECT id FROM templates WHERE id=? AND usuario_id=?', (template_id, user['id']))
        if not await cur.fetchone():
            raise HTTPException(404, 'Template no encontrado o no autorizado')
        await db.execute('DELETE FROM templates WHERE id=?', (template_id,))
        await db.commit()
    return {'status': 'ok'}


# ââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââ
# MAIN
# ââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââââ

if __name__ == '__main__':
    print("""
ââââââââââââââââââââââââââââââââââââââââââââââââ
â   R E V O L V E R   v2   â   Backend        â
â   FastAPI Â· WebSockets Â· SQLite Â· Async      â
ââââââââââââââââââââââââââââââââââââââââââââââââ
""")
    uvicorn.run(
        'revolver_v2:app',
        host='0.0.0.0',
        port=PORT,
        reload=False,
        log_level='info',
    )
